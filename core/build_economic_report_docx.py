from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
ECONOMIC_DIR = PROJECT_ROOT / "output" / "economic_analysis"
OUT_DOCX = ECONOMIC_DIR / "水利经济计算章节.docx"

SCHEME_NAMES = {120.0: "一", 115.0: "二", 108.0: "三", 100.0: "四"}
LEVELS = [120.0, 115.0, 108.0, 100.0]
YEAR_COLS = [f"第{i}年" for i in range(1, 12)]
COMPENSATION_ANNUAL_CHARGE = {120.0: 1328.0, 115.0: 974.7, 108.0: 647.6, 100.0: 442.3}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fmt(value: str | float, digits: int = 2, blank_zero: bool = False) -> str:
    if value == "" or value is None:
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if blank_zero and abs(number) < 1e-12:
        return ""
    if abs(number) < 1e-12:
        number = 0.0
    if digits == 0:
        return f"{number:.0f}"
    text = f"{number:.{digits}f}"
    return text.rstrip("0").rstrip(".") if "." in text else text


def percent(value: str | float, digits: int = 1) -> str:
    if value == "" or value is None:
        return ""
    number = float(value) * 100
    return fmt(number, digits)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size: int = 9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, size=font_size, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)
    run.bold = True


def component_map(rows: list[dict[str, str]]) -> dict[float, dict[str, dict[str, str]]]:
    data: dict[float, dict[str, dict[str, str]]] = {}
    for row in rows:
        level = float(row["正常蓄水位_m"])
        data.setdefault(level, {})[row["项目"]] = row
    return data


def assumption_map(rows: list[dict[str, str]]) -> dict[float, dict[str, dict[str, str]]]:
    data: dict[float, dict[str, dict[str, str]]] = {}
    for row in rows:
        level = float(row["正常蓄水位_m"])
        data.setdefault(level, {})[row["项目"]] = row
    return data


def row_by_name(rows: list[dict[str, str]], name: str) -> dict[str, str]:
    return next(row for row in rows if row["项目"] == name)


def build_document() -> None:
    components = component_map(read_csv(ECONOMIC_DIR / "economic_components.csv"))
    assumptions = assumption_map(read_csv(ECONOMIC_DIR / "economic_assumptions.csv"))
    comparison = {float(row["正常蓄水位_m"]): row for row in read_csv(ECONOMIC_DIR / "economic_comparison.csv")}
    scheme_tables = {
        level: read_csv(ECONOMIC_DIR / "scheme_tables" / f"economic_table_{int(level)}m.csv")
        for level in LEVELS
    }

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("五、经济计算")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)

    add_heading(doc, "5-1 经济计算概述", 1)
    add_paragraph(doc, "水利工程建设方案比较以投资、运行费和综合效益为主要指标。本设计采用年费用最小法，在满足电网电力、电量要求及综合利用要求的前提下，将各方案折算为同一正常运行期年费用口径进行比较。")
    add_paragraph(doc, "从严格意义上看，水利工程各项参数的确定都与经济问题有关。课程设计中时间和资料有限，因此只以正常蓄水位为主要比较对象，说明水利工程建设中经济分析的一般方法。")
    add_paragraph(doc, "投资、运行费和效益是经济分析中的三个主要指标。工程建设追求的目标是净效益最大，当不同方案能够提供相同的电力、电量及其他综合效益时，该目标可转化为费用最小。")
    add_paragraph(doc, "根据电力工程经济分析的常用方法，水电工程方案比较采用年费用最小法。即在各方案同时满足电网电力、电量要求，并满足防洪等综合利用要求的条件下，将各方案的投资、运行费和效益折算为统一年费用，年费用最小者为经济较优方案。")
    add_paragraph(doc, "由于不同正常蓄水位方案的电力、电量效益不同，需以效益最大的方案一为基准，对其他方案配置等效替代火电站及配套煤矿投资。折算率取 r0=0.10，水电站比较期取 50 年。")
    add_paragraph(doc, "替代火电站使用寿命为 25 年，短于水电站比较期。本设计对替代火电站投资采用等额重置处理：先将第一轮火电站投资折算到施工期末，再按 25 年资本回收系数化算为年费用；该年费用在 50 年比较期内连续计入，相当于包含第 25 年后的重置投资。煤矿额外投资及水电工程相关投资仍按 50 年比较期化算。")

    add_heading(doc, "5-2 分投资和补偿投资的计算", 1)
    add_paragraph(doc, "永久性建筑物、机电设备、临时工程及其他费用共同构成水电工程投资。各项投资不是一次完成，而是随施工进度逐年发生，因此需要根据任务书给出的投资流程比例，将总投资分配到施工期各年。")
    add_paragraph(doc, "施工第十一年为施工最后一年，部分设备和工程可折价回收，故投资流程中第十一年含有回收投资率，表现为负值。各方案总投资逐年分配比例见表5-2-1。")
    add_caption(doc, "表5-2-1 各方案总投资逐年分配比例")
    headers = ["年份", "方案一", "方案二", "方案三", "方案四"]
    rows = []
    for i in range(1, 12):
        rows.append([f"第{i}年"] + [percent(row_by_name(scheme_tables[level], "永久、机电、临时投资（%）")[f"第{i}年"]) for level in LEVELS])
    rows.append(["合计", "100%", "100%", "100%", "100%"])
    add_table(doc, headers, rows, [1.0, 1.1, 1.1, 1.1, 1.1])

    add_paragraph(doc, "永久性建筑物投资主要按大坝坝高估算，机电工程投资主要按装机容量估算，其他投资按任务书给定值采用。坝高由调洪计算所得坝顶高程减坝底高程 30 m 得到，永久性建筑物投资和机电设备投资均按相应表格线性插值或必要外推计算。")
    add_caption(doc, "表5-2-2 各方案总投资分项计算表")
    rows = []
    for label, key, unit, digits in [
        ("正常蓄水位", "正常蓄水位_m", "m", 2),
        ("坝顶高程", "坝顶高程", "m", 2),
        ("坝高", "最大坝高", "m", 2),
        ("永久性建筑物投资", "永久性建筑物投资", "万元", 2),
        ("装机容量", "装机容量_万kW", "万kW", 2),
        ("机电工程投资", "机电设备投资", "万元", 2),
        ("其他投资", "其他投资", "万元", 2),
        ("总投资和", "水电工程投资合计", "万元", 2),
    ]:
        values = []
        for level in LEVELS:
            if key in comparison[level]:
                values.append(fmt(comparison[level][key], digits))
            elif key == "正常蓄水位_m":
                values.append(fmt(level, digits))
            else:
                values.append(fmt(components[level][key]["数值"], digits))
        rows.append([label, unit] + values)
    add_table(doc, ["项目", "单位", "方案一", "方案二", "方案三", "方案四"], rows, [1.8, 0.8, 1.15, 1.15, 1.15, 1.15])

    add_paragraph(doc, "除工程投资外，水库淹没处理和移民安置还需计算水库补偿投资。补偿投资从施工第二年开始发生，并在第一台机组发电前一年基本完成。方案一和其余方案采用不同的补偿投资逐年分配比例。")
    add_caption(doc, "表5-2-3 各方案补偿投资逐年分配比例")
    comp1 = row_by_name(scheme_tables[120.0], "水库补偿（%）")
    comp_other = row_by_name(scheme_tables[115.0], "水库补偿（%）")
    rows = [
        ["年份"] + [str(i) for i in range(2, 9)] + ["合计"],
        ["方案一"] + [percent(comp1[f"第{i}年"]) for i in range(2, 9)] + ["100%"],
        ["其余方案"] + [percent(comp_other[f"第{i}年"]) for i in range(2, 9)] + ["100%"],
    ]
    add_table(doc, rows[0], rows[1:], [0.9] + [0.8] * 8)

    add_caption(doc, "表5-2-4 各方案补偿投资逐年分配 单位：万元")
    rows = []
    for level in LEVELS:
        comp = row_by_name(scheme_tables[level], "水库补偿（万元）")
        rows.append([f"方案{SCHEME_NAMES[level]}"] + [fmt(comp[f"第{i}年"], 2, blank_zero=True) for i in range(2, 9)] + [fmt(comparison[level]["水库补偿投资_万元"], 2)])
    add_table(doc, ["方案"] + [str(i) for i in range(2, 9)] + ["合计"], rows, [0.9] + [0.8] * 8)

    add_heading(doc, "5-3 替代投资的计算", 1)
    add_paragraph(doc, "以方案一为基准，其他方案的系统有效容量和多年平均电能不足部分由替代火电站补足。由于火电站厂用电率高于水电站，替代容量按必需容量差值修正计算；替代电能按五强溪多年平均电能差值并考虑凤滩电站影响后计算。")
    add_paragraph(doc, "替代火电站单位投资按 750 元/kW 计，并按 1.1 的容量修正系数折算，故相当于 825 万元/万kW。配套煤矿投资按 0.07 元/kWh 计，并按 1.05 的电能修正系数折算，故相当于 735 万元/亿kWh。")
    add_paragraph(doc, "替代方案的投资包括火电站本身投资和与之配套的煤矿投资。火电站投资安排在第八年至第十一年，煤矿投资安排在第六年至第十年，其投资流程见表5-3-1。")
    add_caption(doc, "表5-3-1 替代方案投资流程分配比例")
    add_table(doc, ["项目", "第六年", "第七年", "第八年", "第九年", "第十年", "第十一年"], [
        ["替代火电站", "", "", "55%", "40%", "3%", "2%"],
        ["煤矿", "16%", "34%", "35%", "10%", "5%", ""],
    ], [1.2] + [0.9] * 6)

    add_caption(doc, "表5-3-2 替代方案投资分配")
    rows = []
    for level in [115.0, 108.0, 100.0]:
        rows.append([
            f"方案{SCHEME_NAMES[level]}",
            fmt(comparison[level]["替代火电容量_万kW"], 2),
            fmt(comparison[level]["替代火电站投资_万元"], 2),
            fmt(comparison[level]["替代火电电能_亿kWh"], 2),
            fmt(comparison[level]["煤矿额外投资_万元"], 2),
        ])
    add_table(doc, ["方案", "替代容量(万kW)", "替代火电站投资(万元)", "替代电能(亿kWh)", "煤矿额外投资(万元)"], rows, [1.0, 1.3, 1.6, 1.3, 1.6])

    add_heading(doc, "5-4 运行费的计算", 1)
    add_paragraph(doc, "水电站正常年运行费包括电站运行费、大修费和水库补偿提成三部分。电站运行费按装机容量和运行费定额计算；大修费按水工建筑物、房屋交通工程和机电设备分别估算；水库补偿提成按任务书给定数值采用。")
    add_caption(doc, "表5-4-1 水电站运行费")
    rows = []
    for item, key, digits in [
        ("装机容量(万kW)", "装机容量_万kW", 2),
        ("水电站运行费及大修补偿合计(万元)", "水电正常年运行费_万元", 2),
    ]:
        rows.append([item] + [fmt(comparison[level][key], digits) for level in LEVELS])
    add_table(doc, ["项目", "方案一", "方案二", "方案三", "方案四"], rows, [2.2, 1.1, 1.1, 1.1, 1.1])

    add_caption(doc, "表5-4-2 各方案水工建筑物大修费")
    rows = []
    rows.append(["坝高(m)"] + [fmt(components[level]["最大坝高"]["数值"], 2) for level in LEVELS])
    rows.append(["大修费(万元)"] + [fmt(assumptions[level]["水工建筑物大修费_万元"]["取值"], 2) for level in LEVELS])
    add_table(doc, ["项目", "方案一", "方案二", "方案三", "方案四"], rows, [1.4, 1.1, 1.1, 1.1, 1.1])

    add_caption(doc, "表5-4-3 各方案补偿提成")
    rows = [
        ["水库补偿(万元)"] + [fmt(comparison[level]["水库补偿投资_万元"], 2) for level in LEVELS],
        ["水库补偿提成(万元)"] + [fmt(COMPENSATION_ANNUAL_CHARGE[level], 2) for level in LEVELS],
    ]
    add_table(doc, ["项目", "方案一", "方案二", "方案三", "方案四"], rows, [1.6, 1.1, 1.1, 1.1, 1.1])

    add_paragraph(doc, "除水电站运行费外，设置替代火电站的方案还需计算火电站运行费和燃料费。火电站运行费按火电站投资的 8% 计算，燃料费按 0.02 元/kWh 计算。")
    add_caption(doc, "表5-4-4 火电站运行费和燃料费")
    rows = []
    for label, key in [("电站投资(万元)", "替代火电站投资_万元"), ("电站运行费(万元)", "火电正常年运行费_万元"), ("替代电能(亿kWh)", "替代火电电能_亿kWh"), ("燃料费(万元)", "火电正常燃料费_万元")]:
        rows.append([label] + [fmt(comparison[level][key], 2, blank_zero=True) for level in LEVELS])
    add_table(doc, ["项目", "方案一", "方案二", "方案三", "方案四"], rows, [1.6, 1.1, 1.1, 1.1, 1.1])

    add_paragraph(doc, "大型水利枢纽通常边施工边投产，从第一台机组发电至电站达到设计效益之间称为初始运行期。为简化计算，各方案初始运行期均取三年，第九、十、十一年发电量分别按正常运行期的 20%、70%、90% 计，第十二年进入正常运行期。")
    add_paragraph(doc, "水电站初期运行费按正常运行费乘以相应发电量比例计算。火电站初期运行费按任务书口径仅计算燃料费，即以正常运行期火电站燃料费乘以相应电量比例。")
    add_paragraph(doc, "五强溪水电站兼有显著防洪效益，不同正常蓄水位对应的汛限水位和防洪库容不同，故各方案防洪效益也不同。防洪效益作为负费用计入经济计算，按多年平均减少拦洪量、减少淹没面积及单位面积效益估算。")
    add_caption(doc, "表5-4-5 各方案防洪效益")
    rows = [
        ["防洪库容(亿m3)"] + [fmt(comparison[level]["防洪库容_亿m3"], 3) for level in LEVELS],
        ["多年平均减少拦洪量(亿m3)"] + [fmt(comparison[level]["多年平均减少拦洪量_亿m3"], 3) for level in LEVELS],
        ["防洪年效益(万元)"] + [fmt(comparison[level]["防洪年效益_万元"], 2) for level in LEVELS],
    ]
    add_table(doc, ["项目", "方案一", "方案二", "方案三", "方案四"], rows, [2.0, 1.1, 1.1, 1.1, 1.1])

    add_heading(doc, "5-5 经济计算总表", 1)
    add_paragraph(doc, "将施工期投资、补偿投资、替代投资、初期运行费及正常运行期运行费、防洪效益汇总，建立各方案资金流程表。第1年至第11年列入施工期和初始运行期实际发生的投资及费用，正常运行期列入正常年运行费、火电站运行费、燃料费及防洪效益。")
    add_paragraph(doc, "第1年至第11年的各年费用先按折算率 10% 折算到第11年末，即施工期末。若第 t 年发生费用 Kt，则折算到施工期末的数值为 Kt(1+r)^(11-t)。各年折算值相加得到施工期末折算总值。")
    add_paragraph(doc, "施工期末折算总值再化算为正常运行期等额年费用。水电工程投资、水库补偿、煤矿投资和初期运行费按 50 年比较期化算；替代火电站投资因使用寿命为 25 年，按 25 年等额重置化算。正常运行期每年直接发生的运行费和防洪效益不再折算为施工期费用，而是作为正常运行期年值直接与化算年费用相加，得到总年费用。")

    for level in LEVELS:
        add_caption(doc, f"表5-5-{SCHEME_NAMES[level]} 方案{SCHEME_NAMES[level]}经济分析表")
        raw_rows = scheme_tables[level]
        headers = ["项目"] + [str(i) for i in range(1, 12)] + ["正常", "折算到施工期末", "化算年费"]
        rows = []
        for row in raw_rows:
            name = row["项目"]
            if name == "火电站投资合计（万元，小计不重复计入总计）":
                name = "火电站投资合计（小计）"
            rows.append([name] + [fmt(row[f"第{i}年"], 2, blank_zero=True) for i in range(1, 12)] + [
                row["正常运行期年值"] if "%" in row["正常运行期年值"] else fmt(row["正常运行期年值"], 2, blank_zero=True),
                fmt(row["折算到施工期末_万元"], 2, blank_zero=True),
                fmt(row["化算到正常运行期年费_万元"], 2, blank_zero=True),
            ])
        add_table(doc, headers, rows, [1.75] + [0.55] * 11 + [0.75, 0.95, 0.85], font_size=7)

    add_heading(doc, "5-6 方案比较", 1)
    add_caption(doc, "表5-6-1 各方案经济比较结果")
    rows = []
    for level in LEVELS:
        row = comparison[level]
        rows.append([
            f"方案{SCHEME_NAMES[level]}",
            fmt(level, 0),
            fmt(row["资本年费用_万元"], 2),
            fmt(row["正常运行期年费用_万元"], 2),
            fmt(row["总年费用_万元"], 2),
            row["是否推荐方案"],
        ])
    add_table(doc, ["方案", "正常蓄水位(m)", "化算到正常运行期年费(万元)", "正常运行期年值(万元)", "总年费用(万元)", "推荐"], rows, [0.9, 1.2, 1.8, 1.5, 1.4, 0.8])
    add_paragraph(doc, "由表可见，方案二正常蓄水位为 115 m 时总年费用最小，为 31249.90 万元/年，故经济比较推荐方案二。")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
